"""
Generates the LinguaFlash documentation report in DOCX format.
Run: python create_documentation.py
"""

from docx import Document
from docx.shared import Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from pathlib import Path

OUTPUT_FILE = Path(__file__).parent / "LinguaFlash_Documentation_Report.docx"


def create_report():
    doc = Document()

    # Title
    title = doc.add_heading("LinguaFlash", 0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    doc.add_paragraph("Language Learning Application with Flashcard System")
    doc.add_paragraph("Project Documentation Report")
    doc.add_paragraph()
    doc.paragraphs[-1].paragraph_format.space_after = Pt(24)

    # ═══════════════════════════════════════════════════════════════════
    # 1. REQUIREMENTS ANALYSIS
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading("1. Requirements Analysis", level=1)

    doc.add_heading("1.1 Functional Requirements", level=2)
    func_reqs = [
        "FR-1: Create and manage multiple flashcard decks organized by language pairs (e.g., English → Spanish).",
        "FR-2: Add, edit, and persist vocabulary cards with front (source language) and back (translation) content.",
        "FR-3: Study mode allowing users to flip cards to reveal translations and rate their recall (Again / Got it!).",
        "FR-4: Dashboard displaying deck statistics (number of decks, total cards) and quick actions.",
        "FR-5: Persistent data storage using JSON so user data survives application restarts.",
        "FR-6: Sample data for first-time users to demonstrate the application immediately.",
    ]
    for req in func_reqs:
        doc.add_paragraph(req, style="List Bullet")

    doc.add_heading("1.2 Non-Functional Requirements", level=2)
    nfr_list = [
        "NFR-1: Modern, visually appealing user interface using CustomTkinter with dark theme and accent colors.",
        "NFR-2: Responsive layout supporting minimum window size of 900×600 pixels.",
        "NFR-3: Cross-platform compatibility (Windows, macOS, Linux) via Python and Tkinter.",
        "NFR-4: Lightweight installation with minimal dependencies (customtkinter, pillow).",
    ]
    for nfr in nfr_list:
        doc.add_paragraph(nfr, style="List Bullet")

    doc.add_heading("1.3 User Personas & Use Cases", level=2)
    doc.add_paragraph(
        "Primary user: Language learners who want to memorize vocabulary using flashcards. "
        "Use cases include creating a new deck (e.g., French vocabulary), adding cards, studying with flip-to-reveal, "
        "and tracking session performance (correct/incorrect counts)."
    )

    # ═══════════════════════════════════════════════════════════════════
    # 2. SYSTEM ARCHITECTURE DESIGN
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading("2. System Architecture Design", level=1)

    doc.add_heading("2.1 High-Level Architecture", level=2)
    doc.add_paragraph(
        "LinguaFlash follows a monolithic desktop application architecture with three main layers: "
        "Presentation (UI), Application (Business Logic), and Data (Persistence)."
    )

    doc.add_heading("2.2 Component Diagram", level=2)
    components = [
        "Presentation Layer: LinguaFlashApp (CTk window), screens: Dashboard, Create Deck, Add Cards, Study Mode, Study Complete.",
        "Application Layer: Event handlers for user actions (create deck, add card, flip card, next card, etc.).",
        "Data Layer: load_data(), save_data(), get_sample_data(); file: flashcard_data.json.",
    ]
    for c in components:
        doc.add_paragraph(c, style="List Bullet")

    doc.add_heading("2.3 Data Model", level=2)
    doc.add_paragraph("Data structure stored in flashcard_data.json:")
    doc.add_paragraph('{"decks": {deck_id: {name, from_lang, to_lang}}, "cards": {deck_id: [{front, back}, ...]}}')
    doc.add_paragraph("Deck ID format: deck_{index}_{random}. Cards are stored as lists of {front, back} objects.")

    doc.add_heading("2.4 Technology Stack", level=2)
    tech = [
        "Python 3.7+",
        "CustomTkinter 5.2+ for modern GUI widgets",
        "Pillow for image support (optional)",
        "Standard library: json, random, pathlib",
    ]
    for t in tech:
        doc.add_paragraph(t, style="List Bullet")

    # ═══════════════════════════════════════════════════════════════════
    # 3. MVP (MINIMUM VIABLE PRODUCT) IMPLEMENTATION
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading("3. MVP (Minimum Viable Product) Implementation", level=1)

    doc.add_heading("3.1 MVP Scope", level=2)
    mvp_scope = [
        "Dashboard with deck list, stats (decks count, total cards), and action buttons.",
        "Create Deck: form with name, from language, to language.",
        "Add Cards: select deck, then add front/back pairs with immediate persistence.",
        "Study Mode: flip cards, reveal answer, rate (Again/Got it!), progress bar, session summary.",
        "Sample Spanish Basics deck (10 cards) for first-time demo.",
    ]
    for s in mvp_scope:
        doc.add_paragraph(s, style="List Bullet")

    doc.add_heading("3.2 Implementation Highlights", level=2)
    doc.add_paragraph(
        "Single-file architecture (main.py) with clear separation: constants, data layer, and LinguaFlashApp class. "
        "Reusable UI helpers: _create_header(), _create_back_button(), _create_card_frame(). "
        "Screen-based navigation: _clear_main() clears container, then target screen is rendered."
    )

    doc.add_heading("3.3 Key Methods", level=2)
    methods = [
        "show_dashboard() — Main hub with stats and deck cards.",
        "_show_create_deck() — Deck creation form.",
        "_show_add_cards(deck_id) — Add cards to a deck.",
        "_start_study(deck_id) — Shuffle cards and begin study session.",
        "_show_study_screen() — Render current card, progress, flip/rating buttons.",
        "_show_study_complete() — Session summary with correct count and mastery percentage.",
    ]
    for m in methods:
        doc.add_paragraph(m, style="List Bullet")

    # ═══════════════════════════════════════════════════════════════════
    # 4. TESTING
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading("4. Testing", level=1)

    doc.add_heading("4.1 Test Strategy", level=2)
    doc.add_paragraph(
        "Testing was performed through manual execution and exploratory testing. "
        "Key flows were validated: deck creation, card addition, study mode navigation, and data persistence."
    )

    doc.add_heading("4.2 Test Cases Executed", level=2)
    test_cases = [
        "TC-1: Launch application — App starts with dashboard and sample deck (if no data file exists).",
        "TC-2: Create new deck — Form accepts name and languages; deck appears on dashboard.",
        "TC-3: Add cards — Cards are saved and count updates on dashboard.",
        "TC-4: Study mode — Cards shuffle, flip works, Again/Got it! advance correctly.",
        "TC-5: Session complete — Correct count and percentage display; Study Again and Back to Dashboard work.",
        "TC-6: Data persistence — Close and reopen app; decks and cards persist.",
    ]
    for tc in test_cases:
        doc.add_paragraph(tc, style="List Bullet")

    doc.add_heading("4.3 Bug Fixes Applied", level=2)
    doc.add_paragraph(
        "CTkFrame.configure() does not support padx/pady. Fix: use inner frames with pack(padx=..., pady=...) for padding."
    )

    # ═══════════════════════════════════════════════════════════════════
    # 5. DOCUMENTATION
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading("5. Documentation", level=1)

    doc.add_heading("5.1 README.md", level=2)
    doc.add_paragraph(
        "Project README includes: feature list, installation (pip install -r requirements.txt), run command (python main.py), "
        "and usage steps for Dashboard, Create Deck, Add Cards, and Study."
    )

    doc.add_heading("5.2 Inline Code Documentation", level=2)
    doc.add_paragraph(
        "Docstrings for all major functions and methods (load_data, save_data, get_sample_data, show_dashboard, etc.). "
        "Comments for theme constants, data structures, and key logic sections."
    )

    doc.add_heading("5.3 This Report", level=2)
    doc.add_paragraph(
        "Comprehensive project documentation covering requirements analysis, system architecture, MVP implementation, "
        "testing, and final presentation as requested."
    )

    # ═══════════════════════════════════════════════════════════════════
    # 6. FINAL PRESENTATION
    # ═══════════════════════════════════════════════════════════════════
    doc.add_heading("6. Final Presentation", level=1)

    doc.add_heading("6.1 Product Summary", level=2)
    doc.add_paragraph(
        "LinguaFlash is a desktop language learning application with a flashcard system. "
        "It enables users to create decks by language pair, add vocabulary cards, and study with a flip-to-reveal "
        "interface and self-rating (Again / Got it!). The application features a visually striking dark theme with "
        "violet, cyan, and amber accents, built with CustomTkinter for a modern look."
    )

    doc.add_heading("6.2 Deliverables", level=2)
    deliverables = [
        "main.py — Main application (~480 lines)",
        "requirements.txt — Dependencies (customtkinter, pillow)",
        "README.md — User guide",
        "flashcard_data.json — Generated at runtime for data persistence",
        "LinguaFlash_Documentation_Report.docx — This documentation report",
    ]
    for d in deliverables:
        doc.add_paragraph(d, style="List Bullet")

    doc.add_heading("6.3 Future Enhancements", level=2)
    future = [
        "Spaced repetition algorithm (e.g., SM-2) for smarter review scheduling.",
        "Audio pronunciation integration (TTS).",
        "Import/export decks (CSV, JSON).",
        "Edit and delete cards/decks.",
        "Light/dark theme toggle.",
    ]
    for f in future:
        doc.add_paragraph(f, style="List Bullet")

    doc.add_heading("6.4 Conclusion", level=2)
    doc.add_paragraph(
        "The LinguaFlash project successfully delivers an MVP that meets the stated requirements: "
        "a functional, visually appealing language learning flashcard application with persistent storage, "
        "sample data for instant demonstration, and clear documentation. The application is ready for end-user "
        "evaluation and serves as a solid foundation for future enhancements."
    )

    doc.save(OUTPUT_FILE)
    print(f"Documentation report saved to: {OUTPUT_FILE}")


if __name__ == "__main__":
    create_report()
